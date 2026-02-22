# import os
# import json
# import re
# from typing import Optional, Dict
# from dotenv import load_dotenv
# from openai import OpenAI
# from pypdf import PdfReader
# from docx import Document

# # Load environment variables
# load_dotenv()

# class JobApplicationAssistant:
#     def __init__(self, api_key: Optional[str] = None):
#         # Initializing for Google Gemini
#         self.client = OpenAI(
#             base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
#             api_key=api_key or os.getenv("GEMINI_API_KEY")
#         )

#     def extract_text_from_file(self, file_path: str) -> str:
#         """Extracts text from PDF or DOCX files."""
#         if not os.path.exists(file_path):
#             raise FileNotFoundError(f"Could not find {file_path}")
        
#         ext = file_path.split('.')[-1].lower()
#         text = ""

#         if ext == 'pdf':
#             reader = PdfReader(file_path)
#             for page in reader.pages:
#                 text += page.extract_text() + "\n"
#         elif ext == 'docx':
#             doc = Document(file_path)
#             for para in doc.paragraphs:
#                 text += para.text + "\n"
#         else:
#             raise ValueError("Unsupported file format. Please use PDF or DOCX.")
            
#         return text

#     def check_experience_gap(self, job_description: str, user_yoe: str) -> Dict:
#         """
#         Analyzes the JD to find minimum experience required and compares with user.
#         """
#         prompt = f"""
#         Analyze the following Job Description text.
        
#         TASK: Extract the minimum years of experience required for this role.
        
#         RULES:
#         1. Return valid JSON only: {{"min_years_required": <float>}}
#         2. If the JD says "3-5 years", return 3.
#         3. If the JD says "5+ years", return 5.
#         4. If NO experience is explicitly mentioned, return 0.
#         5. Do not count "preferred" experience, only "required" or "minimum".
        
#         --- JOB DESCRIPTION ---
#         {job_description}
#         """

#         response = self.client.chat.completions.create(
#             model="gemini-2.5-flash",
#             messages=[{"role": "user", "content": prompt}],
#             temperature=0.0
#         )
        
#         content = response.choices[0].message.content
#         # Clean potential markdown formatting
#         cleaned = re.sub(r'^```json\s*', '', content, flags=re.IGNORECASE)
#         cleaned = re.sub(r'^```\s*', '', cleaned)
#         cleaned = re.sub(r'\s*```$', '', cleaned)
        
#         try:
#             jd_data = json.loads(cleaned)
#             required = float(jd_data.get("min_years_required", 0))
#             user_exp = float(user_yoe)
            
#             return {
#                 "gap_detected": user_exp < required,
#                 "required": required,
#                 "user": user_exp
#             }
#         except:
#             return {"gap_detected": False, "required": 0, "user": 0}

#     def parse_resume_to_json(self, resume_text: str, schema_template: str) -> Dict:
#         """
#         Extracts information from raw resume text and formats it into the provided JSON schema.
#         """
#         prompt = f"""
#         You are an expert HR data parser. 
#         Your task is to extract the candidate's information from the raw resume text and format it EXACTLY according to the JSON schema provided.
        
#         INSTRUCTIONS:
#         1. Return ONLY valid JSON. Do not include any conversational text.
#         2. If a piece of information is missing from the resume, leave the string empty "" or the array empty []. DO NOT make up information.
#         3. Infer 'proficiency_level' (beginner/intermediate/advanced/expert) based on context if possible, otherwise default to "intermediate".
        
#         --- JSON SCHEMA TEMPLATE ---
#         {schema_template}
        
#         --- RAW RESUME TEXT ---
#         {resume_text}
#         """

#         response = self.client.chat.completions.create(
#             model="gemini-2.5-flash", 
#             messages=[
#                 {"role": "system", "content": "You are a precise data extraction AI. Output strictly in JSON format."},
#                 {"role": "user", "content": prompt}
#             ],
#             temperature=0.1
#         )

#         raw_output = response.choices[0].message.content

#         # Robust JSON cleaning
#         cleaned_output = re.sub(r'^```json\s*', '', raw_output, flags=re.IGNORECASE)
#         cleaned_output = re.sub(r'^```\s*', '', cleaned_output)
#         cleaned_output = re.sub(r'\s*```$', '', cleaned_output)
        
#         try:
#             return json.loads(cleaned_output)
#         except json.JSONDecodeError as e:
#             print("Failed to parse JSON from the model. Raw output was:")
#             print(raw_output)
#             raise e

#     def generate_cover_letter(self, job_details: Dict, candidate_profile: Dict) -> str:
#         """Generates a tailored cover letter using the structured JSON profile."""
        
#         prompt = f"""
#         You are an elite Executive Career Coach and Expert Copywriter. Your task is to generate a high-quality, tailored cover letter.

#         ---
#         You must strictly follow the process below.

#         INPUT DATA:

#         [TARGET JOB]
#         Company: {job_details.get('company', 'Unknown')}
#         Role: {job_details.get('role', 'Unknown')}
#         Job Description:
#         {job_details.get('description', '')}

#         [CANDIDATE STRUCTURED DATA - JSON]
#         {json.dumps(candidate_profile, indent=2)}

#         ------------------------------------------------------------
#         PHASE 1 & 2: REASONING & EVIDENCE MAPPING
#         Before writing the letter, you must create a strategy. Write your analysis inside <scratchpad> tags.
#         Inside the <scratchpad>, you must:
#         1. Extract the key technical skills, implied business problems, and seniority expectations from the Job Description.
#         2. Identify quantified achievements from the JSON that directly solve those business problems.
#         3. Select ONE major project to highlight that best aligns with the role.
#         4. If no direct match exists, map out transferable skills.
#         *Do not fabricate tools, metrics, roles, or responsibilities.*

#         ------------------------------------------------------------
#         PHASE 3: COVER LETTER GENERATION
#         After closing the </scratchpad> tag, write the final cover letter.

#         Constraints for the Cover Letter:
#         - 3 to 5 paragraphs maximum.
#         - Standard business letter format.
#         - Tone: Confident, direct, and professional. 
#         - Avoid clichés, generic enthusiasm ("I was thrilled to see..."), and filler text.
#         - Do not restate the resume line-by-line. Focus on problem-solution-impact.
#         - Use quantified impact when available.
#         - No exaggerated claims. 
#         - No placeholders (e.g., do not write [Company Address]. Just skip what you don't know).
#         - You must state the candidate has exactly {candidate_profile.get('personal_info', {}).get('explicit_years_of_experience', 'an unspecified number of')} years of experience. Do not calculate or infer this number from the resume dates.
        
#         Output format:
#         <scratchpad>
#         (Your analysis here)
#         </scratchpad>
#         (The final cover letter here)
#         """

#         response = self.client.chat.completions.create(
#             model="gemini-2.5-flash",
#             messages=[
#                 {"role": "system", "content": "You are an expert cover letter writer."},
#                 {"role": "user", "content": prompt}
#             ],
#             temperature=0.7
#         )

#         full_response = response.choices[0].message.content
        
#         # Strip out the <scratchpad> thinking process using regex
#         final_letter = re.sub(r'<scratchpad>.*?</scratchpad>', '', full_response, flags=re.DOTALL).strip()

#         return final_letter

#     def refine_cover_letter(self, current_letter: str, style_instruction: str) -> str:
#         """Modifies the existing cover letter based on user preference."""
        
#         prompt = f"""
#         You are an expert professional editor. 
#         Your task is to rewrite the provided cover letter based on the following instruction:
        
#         INSTRUCTION: Make the letter {style_instruction}.
        
#         Do not lose the core quantified achievements or the professional tone.
#         Output ONLY the revised cover letter. Do not include any conversational filler or reasoning.
        
#         --- CURRENT COVER LETTER ---
#         {current_letter}
#         """

#         response = self.client.chat.completions.create(
#             model="gemini-2.5-flash",
#             messages=[
#                 {"role": "system", "content": "You are a professional editor."},
#                 {"role": "user", "content": prompt}
#             ],
#             temperature=0.7
#         )

#         return response.choices[0].message.content.strip()


# # ==========================================
# # Main Execution Flow
# # ==========================================
# if __name__ == "__main__":
#     assistant = JobApplicationAssistant()
    
#     PROFILE_FILE = "my_profile.json"

#     # Define the template schema
#     JSON_TEMPLATE = """
#     {
#       "personal_info": { "full_name": "", "email": "", "phone": "", "location": "", "portfolio_url": "", "linkedin_url": "" },
#       "professional_summary": "",
#       "core_competencies": [ { "skill": "", "proficiency_level": "", "years_experience": 0 } ],
#       "work_experience": [ { "company": "", "role": "", "duration": "", "industry": "", "responsibilities": [""], "achievements": [ { "description": "", "impact_metric": "", "tools_used": [""] } ] } ],
#       "major_projects": [ { "name": "", "problem_statement": "", "solution_summary": "", "technologies": [""], "measurable_outcomes": "", "business_impact": "" } ],
#       "education": [ { "degree": "", "institution": "", "year": "", "notable_coursework": [""] } ],
#       "certifications": [""],
#       "leadership_experience": [ { "role": "", "organization": "", "impact": "" } ],
#       "awards": [""],
#       "publication_or_research": [""],
#       "role_alignment_notes": [""]
#     }
#     """

#     # --- STEP 1: RESUME PARSING (Only if JSON doesn't exist) ---
#     if not os.path.exists(PROFILE_FILE):
#         print("No existing profile found. Parsing resume...")
#         resume_path = "resume.pdf" # Replace with your actual resume name
        
#         raw_text = assistant.extract_text_from_file(resume_path)
#         extracted_json = assistant.parse_resume_to_json(raw_text, JSON_TEMPLATE)
#         print("\n--- Quick Clarification ---")
#         explicit_yoe = input("How many years of relevant full-time experience do you have? (e.g., '3', '2.5'): ")
        
#         # Inject the user's explicit answer directly into the JSON
#         extracted_json["personal_info"]["explicit_years_of_experience"] = explicit_yoe
        
#         # Save the structured data
#         with open(PROFILE_FILE, "w") as f:
#             json.dump(extracted_json, f, indent=2)
            
#         print(f"\n[!] SUCCESS: Extracted resume to '{PROFILE_FILE}'.")
#         print("-> Action Required: Open 'my_profile.json' in a text editor.")
#         print("-> Verify the details, fill in any missing gaps (like career goals), and save the file.")
#         input("-> Press ENTER when you are done editing the JSON file...")
#     else:
#         print(f"\nFound existing '{PROFILE_FILE}'. Loading profile...")

#     # --- STEP 2: LOAD VERIFIED JSON ---
#     with open(PROFILE_FILE, "r") as f:
#         user_profile = json.load(f)

#     # --- STEP 3: DYNAMIC JOB DETAILS ---
#     print("\n" + "="*40)
#     print("🎯 NEW JOB APPLICATION")
#     print("="*40)
    
#     company_name = input("Enter Company Name: ").strip()
#     job_role = input("Enter Job Role: ").strip()

#     # Handling the multi-line Job Description safely
#     print("\n📝 Paste the Job Description below.")
#     print("(Press ENTER, type 'DONE' on a new line, and press ENTER again to finish):")
    
#     jd_lines = []
#     while True:
#         line = input()
#         if line.strip().upper() == 'DONE':
#             break
#         jd_lines.append(line)
        
#     job_description = "\n".join(jd_lines)

#     # Pack it into the dictionary
#     job_details = {
#         "company": company_name,
#         "role": job_role,
#         "description": job_description
#     }

#     if not company_name or not job_role or not job_description:
#         print("Error: Missing job details. Exiting...")
#         exit()

#     # --- STEP 4: GENERATE OUTPUTS ---
# # --- STEP 4: EXPERIENCE CHECK & GENERATION ---
#     print(f"\nAnalyzing Job Requirements for {job_details['company']}...")
    
#     # Get user's YoE from the loaded profile
#     user_yoe = user_profile.get('personal_info', {}).get('explicit_years_of_experience', '0')
    
#     # Run the check
#     check_result = assistant.check_experience_gap(job_details['description'], user_yoe)
    
#     if check_result['gap_detected']:
#         print("\n" + "!"*50)
#         print(f"⚠️  WARNING: EXPERIENCE GAP DETECTED")
#         print(f"   - Job Requires: {check_result['required']} years")
#         print(f"   - You Have:     {check_result['user']} years")
#         print("!"*50)
#         proceed = input("\nDo you still want to generate the cover letter? (y/n): ").lower()
#         if proceed != 'y':
#             print("Aborting application.")
#             exit()
#     else:
#         print(f"✅ Experience Check Passed (Job requires {check_result['required']}+ years).")

#     print(f"\nGenerating initial Cover Letter...")
    
#     # ... Continue with generation ...
#     cover_letter = assistant.generate_cover_letter(job_details, user_profile)
#     # (The rest of your code remains the same from here)
#     output_filename = f"Cover_Letter_{job_details['company'].replace(' ', '_')}.txt"
    
#     print("\n" + "="*40)
#     print(cover_letter)
#     print("="*40)
#     # --- STEP 5: INTERACTIVE REFINEMENT ---
#     while True:
#         print("\n--- Refine Cover Letter ---")
#         print("1. Make it more concise and crisp")
#         print("2. Make it more detailed")
#         print("3. Looks good! Save and exit")
        
#         choice = input("Enter your choice (1, 2, or 3): ").strip()
        
#         if choice == '1':
#             print("\nRefining... making it concise and crisp...")
#             cover_letter = assistant.refine_cover_letter(cover_letter, "more concise, crisp, and punchy. Remove fluff")
#             print("\n" + "="*40 + "\n" + cover_letter + "\n" + "="*40)
            
#         elif choice == '2':
#             print("\nRefining... making it more detailed...")
#             cover_letter = assistant.refine_cover_letter(cover_letter, "more detailed. Expand on the problem-solving aspects while keeping it professional")
#             print("\n" + "="*40 + "\n" + cover_letter + "\n" + "="*40)
            
#         elif choice == '3':
#             with open(output_filename, "w") as f:
#                 f.write(cover_letter)
#             print(f"\n[!] Final cover letter saved to '{output_filename}'")
#             break
            
#         else:
#             print("Invalid choice. Please enter 1, 2, or 3.")










import os
import json
import re
from typing import Optional, Dict
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document

# Load environment variables
load_dotenv()

class JobApplicationAssistant:
    def __init__(self, api_key: Optional[str] = None):
        # Initializing for Google Gemini
        self.client = OpenAI(
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/",
            api_key=api_key or os.getenv("GEMINI_API_KEY")
        )

    def extract_text_from_file(self, file_path: str) -> str:
        """Extracts text from PDF or DOCX files."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Could not find {file_path}")
        
        ext = file_path.split('.')[-1].lower()
        text = ""

        if ext == 'pdf':
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif ext == 'docx':
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            raise ValueError("Unsupported file format. Please use PDF or DOCX.")
            
        return text

    def parse_resume_to_json(self, resume_text: str, schema_template: str) -> Dict:
        """Extracts information from raw resume text and formats it into JSON."""
        prompt = f"""
        You are an expert HR data parser. 
        Your task is to extract the candidate's information from the raw resume text and format it EXACTLY according to the JSON schema provided.
        
        INSTRUCTIONS:
        1. Return ONLY valid JSON.
        2. If a piece of information is missing, leave it empty.
        3. Infer 'proficiency_level' based on context.
        
        --- JSON SCHEMA TEMPLATE ---
        {schema_template}
        
        --- RAW RESUME TEXT ---
        {resume_text}
        """

        response = self.client.chat.completions.create(
            model="gemini-2.5-flash", 
            messages=[
                {"role": "system", "content": "You are a precise data extraction AI. Output strictly in JSON format."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )

        raw_output = response.choices[0].message.content

        # Robust JSON cleaning
        cleaned_output = re.sub(r'^```json\s*', '', raw_output, flags=re.IGNORECASE)
        cleaned_output = re.sub(r'^```\s*', '', cleaned_output)
        cleaned_output = re.sub(r'\s*```$', '', cleaned_output)
        
        try:
            return json.loads(cleaned_output)
        except json.JSONDecodeError as e:
            print("Failed to parse JSON. Raw output was:", raw_output)
            raise e

    def check_experience_gap(self, job_description: str, user_yoe: str) -> Dict:
        """Analyzes the JD to find minimum experience required and compares with user."""
        prompt = f"""
        Analyze the following Job Description text.
        
        TASK: Extract the minimum years of experience required for this role.
        
        RULES:
        1. Return valid JSON only: {{"min_years_required": <float>}}
        2. If the JD says "3-5 years", return 3.
        3. If the JD says "5+ years", return 5.
        4. If NO experience is explicitly mentioned, return 0.
        
        --- JOB DESCRIPTION ---
        {job_description}
        """

        response = self.client.chat.completions.create(
            model="gemini-2.5-flash",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0
        )
        
        content = response.choices[0].message.content
        cleaned = re.sub(r'^```json\s*', '', content, flags=re.IGNORECASE)
        cleaned = re.sub(r'^```\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)
        
        try:
            jd_data = json.loads(cleaned)
            required = float(jd_data.get("min_years_required", 0))
            user_exp = float(user_yoe) if user_yoe else 0.0
            
            return {
                "gap_detected": user_exp < required,
                "required": required,
                "user": user_exp
            }
        except:
            return {"gap_detected": False, "required": 0, "user": 0}

    def generate_cover_letter(self, job_details: Dict, candidate_profile: Dict) -> str:
        """Generates a tailored cover letter using the structured JSON profile."""
        
        prompt = f"""
        You are an elite Executive Career Coach and Expert Copywriter. Your task is to generate a high-quality, tailored cover letter.

        ---
        You must strictly follow the process below.

        INPUT DATA:

        [TARGET JOB]
        Company: {job_details.get('company', 'Unknown')}
        Role: {job_details.get('role', 'Unknown')}
        Job Description:
        {job_details.get('description', '')}

        [CANDIDATE STRUCTURED DATA - JSON]
        {json.dumps(candidate_profile, indent=2)}

        ------------------------------------------------------------
        PHASE 1 & 2: REASONING & EVIDENCE MAPPING
        Before writing the letter, you must create a strategy. Write your analysis inside <scratchpad> tags.
        Inside the <scratchpad>, you must:
        1. Extract the key technical skills, implied business problems, and seniority expectations from the Job Description.
        2. Identify quantified achievements from the JSON that directly solve those business problems.
        3. Select ONE major project to highlight that best aligns with the role.
        4. If no direct match exists, map out transferable skills.
        *Do not fabricate tools, metrics, roles, or responsibilities.*

        ------------------------------------------------------------
        PHASE 3: COVER LETTER GENERATION
        After closing the </scratchpad> tag, write the final cover letter.

        Constraints for the Cover Letter:
        - 3 to 5 paragraphs maximum.
        - Standard business letter format.
        - Tone: Confident, direct, and professional. 
        - Avoid clichés, generic enthusiasm ("I was thrilled to see..."), and filler text.
        - Do not restate the resume line-by-line. Focus on problem-solution-impact.
        - Use quantified impact when available.
        - No exaggerated claims. 
        - No placeholders (e.g., do not write [Company Address]. Just skip what you don't know).
        - You must state the candidate has exactly {candidate_profile.get('personal_info', {}).get('explicit_years_of_experience', 'an unspecified number of')} years of experience. Do not calculate or infer this number from the resume dates.
        
        Output format:
        <scratchpad>
        (Your analysis here)
        </scratchpad>
        (The final cover letter here)
        """

        response = self.client.chat.completions.create(
            model="gemini-2.5-flash",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )

        full_response = response.choices[0].message.content
        final_letter = re.sub(r'<scratchpad>.*?</scratchpad>', '', full_response, flags=re.DOTALL).strip()
        return final_letter

    def refine_cover_letter(self, current_letter: str, instruction: str, job_details: Dict, candidate_profile: Dict) -> str:
        """
        Modifies the existing cover letter based on user preference, 
        using the full context of the job and profile.
        """
        
        prompt = f"""
        You are an expert professional editor. 
        Your task is to rewrite the provided cover letter based on the user's specific instruction.
        
        --- CONTEXT ---
        Target Role: {job_details.get('role')} at {job_details.get('company')}
        Candidate Profile (Reference for specific skills/projects): 
        {json.dumps(candidate_profile, indent=2)}
        
        --- CURRENT COVER LETTER ---
        {current_letter}
        
        --- USER INSTRUCTION ---
        {instruction}
        
        GUIDELINES:
        1. Implement the user's instruction precisely.
        2. If the user asks to add specific details (like a project or skill), look for them in the Candidate Profile above.
        3. Keep the tone professional unless asked otherwise.
        4. Output ONLY the revised cover letter. No conversational filler.
        """

        response = self.client.chat.completions.create(
            model="gemini-2.5-flash",
            messages=[
                {"role": "system", "content": "You are a professional editor."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )

        return response.choices[0].message.content.strip()


# ==========================================
# Main Execution Flow
# ==========================================
if __name__ == "__main__":
    assistant = JobApplicationAssistant()
    
    PROFILE_FILE = "my_profile.json"

    # Define the template schema
    JSON_TEMPLATE = """
    {
      "personal_info": { "full_name": "", "email": "", "phone": "", "location": "", "portfolio_url": "", "linkedin_url": "" },
      "professional_summary": "",
      "core_competencies": [ { "skill": "", "proficiency_level": "", "years_experience": 0 } ],
      "work_experience": [ { "company": "", "role": "", "duration": "", "industry": "", "responsibilities": [""], "achievements": [ { "description": "", "impact_metric": "", "tools_used": [""] } ] } ],
      "major_projects": [ { "name": "", "problem_statement": "", "solution_summary": "", "technologies": [""], "measurable_outcomes": "", "business_impact": "" } ],
      "education": [ { "degree": "", "institution": "", "year": "", "notable_coursework": [""] } ],
      "certifications": [""],
      "leadership_experience": [ { "role": "", "organization": "", "impact": "" } ],
      "awards": [""],
      "publication_or_research": [""],
      "role_alignment_notes": [""]
    }
    """

    # --- STEP 1: RESUME PARSING ---
    if not os.path.exists(PROFILE_FILE):
        print("No existing profile found. Parsing resume...")
        resume_path = "resume.pdf" 
        
        raw_text = assistant.extract_text_from_file(resume_path)
        extracted_json = assistant.parse_resume_to_json(raw_text, JSON_TEMPLATE)
        print("\n--- Quick Clarification ---")
        explicit_yoe = input("How many years of relevant full-time experience do you have? (e.g., '3', '2.5'): ")
        
        extracted_json["personal_info"]["explicit_years_of_experience"] = explicit_yoe
        
        with open(PROFILE_FILE, "w") as f:
            json.dump(extracted_json, f, indent=2)
            
        print(f"\n[!] SUCCESS: Extracted resume to '{PROFILE_FILE}'.")
        print("-> Action Required: Open 'my_profile.json' in a text editor.")
        print("-> Verify the details and save the file.")
        input("-> Press ENTER when you are done editing the JSON file...")
    else:
        print(f"\nFound existing '{PROFILE_FILE}'. Loading profile...")

    # --- STEP 2: LOAD VERIFIED JSON ---
    with open(PROFILE_FILE, "r") as f:
        user_profile = json.load(f)

    # --- STEP 3: DYNAMIC JOB DETAILS ---
    print("\n" + "="*40)
    print("🎯 NEW JOB APPLICATION")
    print("="*40)
    
    company_name = input("Enter Company Name: ").strip()
    job_role = input("Enter Job Role: ").strip()

    print("\n📝 Paste the Job Description below.")
    print("(Press ENTER, type 'DONE' on a new line, and press ENTER again to finish):")
    
    jd_lines = []
    while True:
        line = input()
        if line.strip().upper() == 'DONE':
            break
        jd_lines.append(line)
        
    job_description = "\n".join(jd_lines)

    job_details = {
        "company": company_name,
        "role": job_role,
        "description": job_description
    }

    if not company_name or not job_role or not job_description:
        print("Error: Missing job details. Exiting...")
        exit()

    # --- STEP 4: EXPERIENCE CHECK ---
    print(f"\nAnalyzing Job Requirements for {job_details['company']}...")
    user_yoe = user_profile.get('personal_info', {}).get('explicit_years_of_experience', '0')
    check_result = assistant.check_experience_gap(job_details['description'], user_yoe)
    
    if check_result['gap_detected']:
        print("\n" + "!"*50)
        print(f"⚠️  WARNING: EXPERIENCE GAP DETECTED")
        print(f"   - Job Requires: {check_result['required']} years")
        print(f"   - You Have:     {check_result['user']} years")
        print("!"*50)
        proceed = input("\nDo you still want to generate the cover letter? (y/n): ").lower()
        if proceed != 'y':
            print("Aborting application.")
            exit()
    else:
        print(f"✅ Experience Check Passed (Job requires {check_result['required']}+ years).")

    # --- STEP 5: GENERATION & REFINEMENT ---
    print(f"\nGenerating initial Cover Letter...")
    cover_letter = assistant.generate_cover_letter(job_details, user_profile)
    output_filename = f"Cover_Letter_{job_details['company'].replace(' ', '_')}.txt"
    
    print("\n" + "="*40)
    print(cover_letter)
    print("="*40)

    while True:
        print("\n--- Refine Cover Letter ---")
        print("1. Make it more concise and crisp")
        print("2. Make it more detailed")
        print("3. Give specific instructions (e.g., 'Mention my AWS cert', 'Make it Friendlier')")
        print("4. Looks good! Save and exit")
        
        choice = input("Enter your choice (1-4): ").strip()
        
        instruction = ""
        if choice == '1':
            instruction = "Make it more concise, crisp, and punchy. Remove fluff."
        elif choice == '2':
            instruction = "Make it more detailed. Expand on problem-solving aspects professionally."
        elif choice == '3':
            instruction = input("Enter your specific instruction: ")
        elif choice == '4':
            with open(output_filename, "w") as f:
                f.write(cover_letter)
            print(f"\n[!] Final cover letter saved to '{output_filename}'")
            break
        else:
            print("Invalid choice. Please enter 1, 2, 3, or 4.")
            continue

        print("\nRefining...")
        # Now passing the full context (job details + profile) to the refiner
        cover_letter = assistant.refine_cover_letter(cover_letter, instruction, job_details, user_profile)
        print("\n" + "="*40 + "\n" + cover_letter + "\n" + "="*40)